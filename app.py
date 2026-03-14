import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- CẤU HÌNH ---
if "GEMINI_API_KEY" in st.secrets:
    MY_API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=MY_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
else:
    st.error("Chưa cấu hình API Key trong phần Secrets của Streamlit!")
    st.stop() # Dừng app nếu không có key

        
        mon = st.text_input("1. Tên môn học:", "Lịch sử")
        noidung = st.text_area("2. Dán nội dung bài học vào đây:", height=250)
        
        if st.button("🔥 BẮT ĐẦU TẠO ĐỀ"):
            if not noidung:
                st.warning("Bạn chưa dán nội dung bài học kìa!")
            else:
                with st.spinner("AI đang soạn đề, đợi tí nhé..."):
                    res = model.generate_content(f"Tạo 10 câu trắc nghiệm từ nội dung này: {noidung}")
                    st.markdown(res.text)
                    
                    doc = Document()
                    doc.add_heading(f'ĐỀ THI MÔN: {mon.upper()}', 0)
                    doc.add_paragraph(res.text)
                    bio = BytesIO()
                    doc.save(bio)
                    st.download_button("📥 TẢI FILE WORD", bio.getvalue(), f"De_{mon}.docx")
    except Exception as e:
        st.error(f"Lỗi hệ thống: {e}")
        
